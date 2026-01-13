import logging
from docx import Document
from .base import BaseExtractor
from .models import ExtractResult

logger = logging.getLogger(__name__)

class DocxExtractor(BaseExtractor):
    def extract(self, path: str) -> ExtractResult:
        try:
            doc = Document(path)
            text = []
            
            # 1. Paragraphs
            for para in doc.paragraphs:
                text.append(para.text)
                
            # 2. Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
            
            # 3. Embedded Images (OCR)
            # Check config
            features_cfg = self.config.get("features", {})
            extraction_cfg = features_cfg.get("extraction", {})
            if not extraction_cfg:
                extraction_cfg = self.config.get("extraction", {})
            
            if extraction_cfg.get("images", False) and extraction_cfg.get("ocr", False):
                 try:
                     from app.core.extractors.image import ImageExtractor
                     img_extractor = ImageExtractor(self.config)
                     import tempfile
                     import os
                     
                     # Iterate relationships safely
                     rels = list(doc.part.rels.values())
                     
                     for i, rel in enumerate(rels):
                        try:
                             content_type = "UNKNOWN"
                             try:
                                 content_type = getattr(rel.target_part, "content_type", "N/A")
                             except Exception:
                                 # Some relationships might be internal or broken
                                 continue

                             if "image" in str(getattr(rel, "target_ref", "")).lower() or (str(content_type).startswith("image/")):
                                 try:
                                     image_data = rel.target_part.blob
                                     # Guess extension
                                     ext = ".png"
                                     if "jpeg" in str(content_type): ext = ".jpg"
                                     elif "png" in str(content_type): ext = ".png"
                                 
                                     with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                                         tmp.write(image_data)
                                         tmp_path = tmp.name
                                         
                                     try:
                                         res = img_extractor.extract(tmp_path)
                                         
                                         if res.content:
                                              text.append(f"\n[IMG-OCR: {res.content}]")
                                         else:
                                              # Use metadata to distinguish empty vs error
                                              if res.metadata.get("source") == "ocr_empty":
                                                  text.append(f"\n[IMG-OCR: <NO TEXT DETECTED>]")
                                     finally:
                                          if os.path.exists(tmp_path):
                                              os.unlink(tmp_path)
                                 except Exception as e:
                                     logger.warning(f"DocxExtractor Image Error (Relationship {i}): {e}") 
                        except Exception as loop_e:
                             logger.warning(f"DocxExtractor Rel Loop Error (Index {i}): {loop_e}")

                 except Exception as e:
                     logger.warning(f"DocxExtractor Init Error: {e}")

            return ExtractResult(content="\n".join(text), metadata={"source": "text"})
        except Exception as e:
            return ExtractResult(content=None, error=str(e), metadata={"source": "error"})
